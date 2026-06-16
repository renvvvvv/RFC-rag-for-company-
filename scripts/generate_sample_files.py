"""Generate sample documents for knowledge base demonstration."""
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from PIL import Image, ImageDraw, ImageFont

SAMPLES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "samples")
os.makedirs(SAMPLES_DIR, exist_ok=True)


def write_markdown(filename: str, title: str, content: str) -> None:
    path = os.path.join(SAMPLES_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"# {title}\n\n")
        f.write(f"> 生成时间：{datetime.now().strftime('%Y-%m-%d')}\n\n")
        f.write(content)
    print(f"Generated: {path}")


def write_txt(filename: str, title: str, content: str) -> None:
    path = os.path.join(SAMPLES_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"{title}\n")
        f.write("=" * 40 + "\n\n")
        f.write(f"生成时间：{datetime.now().strftime('%Y-%m-%d')}\n\n")
        f.write(content)
    print(f"Generated: {path}")


def generate_excel_finance() -> None:
    path = os.path.join(SAMPLES_DIR, "04-财务数据样例.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "2025年度财务概览"

    headers = ["季度", "营业收入(万元)", "营业成本(万元)", "净利润(万元)", "同比增长"]
    data = [
        ["Q1", 12500, 8200, 2300, "12%"],
        ["Q2", 14800, 9600, 2900, "18%"],
        ["Q3", 16200, 10100, 3400, "22%"],
        ["Q4", 18900, 11500, 4100, "25%"],
    ]

    header_fill = PatternFill(start_color="0F172A", end_color="0F172A", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    for r_idx, row in enumerate(data, 2):
        for c_idx, value in enumerate(row, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center" if c_idx > 1 else "left")

    ws.column_dimensions["A"].width = 12
    for col in ["B", "C", "D", "E"]:
        ws.column_dimensions[col].width = 20

    # Add summary sheet
    ws2 = wb.create_sheet("部门预算")
    ws2.append(["部门", "预算(万元)", "已使用(万元)", "剩余(万元)"])
    budget_data = [
        ["研发部", 3500, 2100, 1400],
        ["市场部", 2000, 1500, 500],
        ["销售部", 2800, 1900, 900],
        ["运营部", 1500, 900, 600],
    ]
    for row in budget_data:
        ws2.append(row)

    wb.save(path)
    print(f"Generated: {path}")


def generate_excel_project() -> None:
    path = os.path.join(SAMPLES_DIR, "08-项目计划表.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "RAG项目里程碑"

    headers = ["阶段", "任务", "负责人", "开始日期", "结束日期", "状态"]
    data = [
        ["第一阶段", "需求调研与方案设计", "张三", "2025-01-01", "2025-02-15", "已完成"],
        ["第二阶段", "知识库搭建与数据接入", "李四", "2025-02-16", "2025-04-30", "已完成"],
        ["第三阶段", "RAG引擎开发与测试", "王五", "2025-05-01", "2025-07-31", "进行中"],
        ["第四阶段", "权限体系与安全加固", "赵六", "2025-08-01", "2025-09-15", "未开始"],
        ["第五阶段", "上线部署与运维监控", "钱七", "2025-09-16", "2025-10-31", "未开始"],
    ]

    ws.append(headers)
    for row in data:
        ws.append(row)

    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=len(headers)):
        for cell in row:
            cell.border = Border(
                left=Side(style="thin"), right=Side(style="thin"),
                top=Side(style="thin"), bottom=Side(style="thin")
            )
            if cell.row == 1:
                cell.fill = PatternFill(start_color="E57035", end_color="E57035", fill_type="solid")
                cell.font = Font(color="FFFFFF", bold=True)

    for col in ["A", "B", "C", "D", "E", "F"]:
        ws.column_dimensions[col].width = 20

    wb.save(path)
    print(f"Generated: {path}")


def generate_word_doc() -> None:
    path = os.path.join(SAMPLES_DIR, "05-技术白皮书.docx")
    doc = Document()

    title = doc.add_heading("企业级私有化多模态 RAG 技术白皮书", level=0)
    title.alignment = 1
    run = title.runs[0]
    run.font.color.rgb = RGBColor(15, 23, 42)
    run.font.size = Pt(22)

    doc.add_paragraph(f"版本：v1.0 | 日期：{datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()

    doc.add_heading("1. 引言", level=1)
    doc.add_paragraph(
        "随着大语言模型（LLM）在企业场景中的广泛应用，如何让模型安全、准确地利用企业私有知识，"
        "成为企业智能化转型的关键挑战。通用 LLM 存在数据隐私风险、知识更新滞后、幻觉等问题，"
        "而检索增强生成（RAG）技术通过将外部知识库与生成模型结合，有效缓解了上述问题。"
    )

    doc.add_heading("2. 系统架构", level=1)
    doc.add_paragraph(
        "本系统采用模块化架构，主要包括：数据接入层、多模态解析层、向量索引层、检索引擎层、"
        "权限控制层、生成服务层和网关接入层。"
    )

    doc.add_heading("3. 关键技术", level=1)
    items = [
        "混合检索：结合向量相似度、BM25 和重排序模型，提升召回准确率。",
        "多模态理解：支持文本、表格、图片、音频、视频等多种数据类型的解析与索引。",
        "细粒度权限：从身份、知识库、文档、字段到标签五级权限穿透。",
        "来源追溯：每条生成结果均附带引用来源，支持审计与验证。",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. 部署方式", level=1)
    doc.add_paragraph(
        "系统支持 Docker Compose 一键私有化部署，所有数据存储在企业内部，满足金融、政务、"
        "制造等行业对数据安全和合规的严格要求。"
    )

    doc.save(path)
    print(f"Generated: {path}")


def generate_pdf() -> None:
    path = os.path.join(SAMPLES_DIR, "03-产品功能说明.pdf")
    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=20,
        textColor=colors.HexColor("#0F172A"),
        spaceAfter=20,
        alignment=1,
    )

    story = []
    story.append(Paragraph("企业级私有化多模态 RAG 产品功能说明", title_style))
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(f"<para align=center>生成日期：{datetime.now().strftime('%Y-%m-%d')}</para>", styles["Normal"]))
    story.append(Spacer(1, 1*cm))

    sections = [
        ("一、知识库管理", "支持多知识库创建与管理，每个知识库可独立配置权限、模态和索引策略。"),
        ("二、文档上传", "支持 PDF、Word、Excel、图片、视频、音频、Markdown、网页链接等多种格式上传。"),
        ("三、智能检索", "提供向量检索、关键词检索、混合检索和重排序，支持多知识库联合查询。"),
        ("四、对话问答", "基于检索结果进行流式生成，支持多轮对话、历史消息、引用来源展示。"),
        ("五、权限控制", "五级权限模型，确保用户只能访问被授权的知识和内容。"),
        ("六、评测工作台", "内置 Recall、MRR、NDCG、Faithfulness 等指标，持续优化 RAG 效果。"),
        ("七、可观测性", "集成 Prometheus 和 Grafana，提供全方位的监控和告警能力。"),
    ]

    for title, content in sections:
        story.append(Paragraph(title, styles["Heading2"]))
        story.append(Paragraph(content, styles["Normal"]))
        story.append(Spacer(1, 0.3*cm))

    feature_data = [
        ["功能模块", "核心能力", "适用场景"],
        ["知识库", "多库隔离、权限绑定", "部门级知识管理"],
        ["上传中心", "多模态解析、批量上传", "文档数字化"],
        ["检索控制台", "自然语言问答、来源追溯", "智能客服、知识助手"],
        ["权限管理", "五级权限、敏感词拦截", "安全合规"],
    ]
    table = Table(feature_data, colWidths=[4*cm, 5*cm, 5*cm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#F8FAFC")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))
    story.append(Spacer(1, 0.5*cm))
    story.append(table)

    doc.build(story)
    print(f"Generated: {path}")


def generate_org_chart() -> None:
    path = os.path.join(SAMPLES_DIR, "06-组织架构图.png")
    width, height = 800, 600
    img = Image.new("RGB", (width, height), "#F8FAFC")
    draw = ImageDraw.Draw(img)

    # Try to load a font, fallback to default
    try:
        font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 24)
        font_node = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 16)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
    except Exception:
        font_title = ImageFont.load_default()
        font_node = ImageFont.load_default()
        font_small = ImageFont.load_default()

    # Title
    draw.text((width // 2, 30), "企业 RAG 项目组织架构", fill="#0F172A", font=font_title, anchor="mt")

    # Box dimensions
    box_w, box_h = 140, 60
    nodes = [
        ("项目总监", width // 2, 100, "#0F172A"),
        ("产品经理", width // 2 - 200, 200, "#E57035"),
        ("技术负责人", width // 2, 200, "#E57035"),
        ("安全合规官", width // 2 + 200, 200, "#E57035"),
        ("知识库运营", width // 2 - 250, 320, "#334155"),
        ("算法工程师", width // 2 - 80, 320, "#334155"),
        ("后端工程师", width // 2 + 90, 320, "#334155"),
        ("前端工程师", width // 2 + 260, 320, "#334155"),
        ("数据标注", width // 2 - 150, 440, "#64748B"),
        ("权限管理", width // 2 + 150, 440, "#64748B"),
    ]

    # Draw connections
    def draw_box(x, y, text, color):
        x1, y1 = x - box_w // 2, y - box_h // 2
        x2, y2 = x1 + box_w, y1 + box_h
        draw.rounded_rectangle([x1, y1, x2, y2], radius=10, fill=color, outline="#CBD5E1", width=2)
        draw.text((x, y), text, fill="white" if color in ["#0F172A", "#E57035", "#334155"] else "#1F2937", font=font_node, anchor="mm")

    # Connections
    connections = [
        ((width // 2, 130), (width // 2 - 200, 170)),
        ((width // 2, 130), (width // 2, 170)),
        ((width // 2, 130), (width // 2 + 200, 170)),
        ((width // 2 - 200, 230), (width // 2 - 250, 290)),
        ((width // 2, 230), (width // 2 - 80, 290)),
        ((width // 2, 230), (width // 2 + 90, 290)),
        ((width // 2, 230), (width // 2 + 260, 290)),
        ((width // 2 - 200, 230), (width // 2 - 150, 410)),
        ((width // 2 + 200, 230), (width // 2 + 150, 410)),
    ]
    for (x1, y1), (x2, y2) in connections:
        draw.line([(x1, y1), (x2, y2)], fill="#94A3B8", width=2)

    for text, x, y, color in nodes:
        draw_box(x, y, text, color)

    draw.text((width // 2, 540), "注：本图仅供演示多模态 RAG 的图片解析能力", fill="#6B7280", font=font_small, anchor="mt")

    img.save(path)
    print(f"Generated: {path}")


def main():
    product_intro = """\
## 产品定位

企业级私有化多模态 RAG 系统，面向金融、制造、政务、能源等行业，提供数据不出域、权限可管控、来源可追溯、多模态可理解的企业知识检索与生成平台。

## 核心能力

- **多模态接入**：支持文档、表格、图片、视频、音频、网页链接等多种数据类型。
- **统一检索**：向量检索 + BM25 + 重排序，召回高相关片段。
- **权限管控**：身份、知识库、文档、字段、标签五级权限穿透。
- **可信生成**：每条回答附带引用来源，支持审计与验证。
- **私有化部署**：Docker Compose 一键部署，数据完全存储在企业内部。

## 典型应用场景

1. 企业内部知识库问答助手
2. 合同、标书、技术文档智能检索
3. 培训视频、会议纪要的语义搜索
4. 产品手册、FAQ 的自动化客服

## 联系我们

如需了解更多产品信息或申请试用，请联系企业解决方案团队。\
"""

    faq = """\
## 常见问题

### Q1：系统支持哪些文件格式？

系统支持 PDF、Word（doc/docx）、Excel（xls/xlsx）、图片（png/jpg/gif）、视频（mp4/avi/mov）、音频（mp3/wav）、Markdown、TXT 以及网页链接。

### Q2：数据是否会离开企业内部？

不会。系统采用私有化部署，所有数据存储在企业内部的服务器上，向量数据库、对象存储、关系数据库均由企业自行管理。

### Q3：如何保证不同用户看到不同的内容？

系统提供五级权限模型：身份层、知识库层、文档层、字段层和标签层。检索结果会经过统一权限服务过滤，确保用户只能看到被授权的内容。

### Q4：生成的答案是否可信？

系统为每条生成答案提供引用来源，包括文档 ID、Chunk ID、相关度分数和模态信息，用户可以追溯到原始内容，降低大模型幻觉风险。

### Q5：是否支持与现有系统集成？

支持。系统提供基于 Kong 网关的标准 RESTful API，支持 JWT 和 API Key 鉴权，可方便地接入 OA、IM、ERP、BI 等现有业务系统。\
"""

    txt_content = """\
企业级私有化多模态 RAG 系统

产品定位：面向金融、制造、政务、能源等行业，提供数据不出域、权限可管控、来源可追溯、多模态可理解的企业知识检索与生成平台。

核心能力：
1. 多模态接入：支持文档、表格、图片、视频、音频、网页链接等。
2. 统一检索：向量检索 + BM25 + 重排序。
3. 权限管控：身份、知识库、文档、字段、标签五级权限穿透。
4. 可信生成：每条回答附带引用来源。
5. 私有化部署：Docker Compose 一键部署。

典型应用场景：
- 企业内部知识库问答助手
- 合同、标书、技术文档智能检索
- 培训视频、会议纪要的语义搜索
- 产品手册、FAQ 的自动化客服
"""

    write_markdown("01-企业RAG产品介绍.md", "企业级私有化多模态 RAG 产品介绍", product_intro)
    write_txt("02-企业RAG产品介绍.txt", "企业级私有化多模态 RAG 产品介绍", txt_content)
    write_markdown("07-客户服务FAQ.md", "客户服务常见问题", faq)

    generate_pdf()
    generate_word_doc()
    generate_excel_finance()
    generate_excel_project()
    generate_org_chart()

    # README
    readme_path = os.path.join(SAMPLES_DIR, "README.md")
    with open(readme_path, "w", encoding="utf-8") as f:
        f.write("# 知识库样例文件\n\n")
        f.write("本目录包含用于演示企业级私有化多模态 RAG 系统能力的样例文件。\n\n")
        f.write("## 文件清单\n\n")
        f.write("| 文件名 | 类型 | 用途 |\n")
        f.write("|---|---|---|\n")
        f.write("| 01-企业RAG产品介绍.md | Markdown | 产品介绍，测试文本解析 |\n")
        f.write("| 02-企业RAG产品介绍.txt | 文本 | 纯文本版本，测试 TXT 解析 |\n")
        f.write("| 03-产品功能说明.pdf | PDF | 产品功能说明，测试 PDF 解析 |\n")
        f.write("| 04-财务数据样例.xlsx | Excel | 财务数据，测试表格解析 |\n")
        f.write("| 05-技术白皮书.docx | Word | 技术文档，测试 Word 解析 |\n")
        f.write("| 06-组织架构图.png | 图片 | 组织架构图，测试图片 OCR |\n")
        f.write("| 07-客户服务FAQ.md | Markdown | FAQ 内容，测试问答召回 |\n")
        f.write("| 08-项目计划表.xlsx | Excel | 项目里程碑，测试表格解析 |\n")
        f.write("\n## 使用方式\n\n")
        f.write("1. 进入系统后访问\"上传中心\"页面。\n")
        f.write("2. 选择目标知识库。\n")
        f.write("3. 拖拽或点击上传本目录中的文件。\n")
        f.write("4. 文件解析索引完成后，在\"检索控制台\"进行问答测试。\n")
    print(f"Generated: {readme_path}")


if __name__ == "__main__":
    main()
